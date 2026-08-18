from __future__ import annotations

import asyncio
import hashlib
import shutil
import sys
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock

import pytest

from app.models import ContentParseOptions, StoredSource
from app.models.content_result import (
    AssetKind,
    AssetLocation,
    AssetRole,
    SourceKind,
    UnitType,
    VisualAnalysis,
    VisualClassification,
)
from app.models.parse_result import PageSourceKind
from app.parsers.base import ParserCancelledError, ParserUnavailableError
from app.security.file_validation import DOCX_MIME, PPTX_MIME
from app.services.evidence_service import PageEvidence
from app.services.multimodal_service import MultimodalService
from app.services.storage_service import StorageService


def service(tmp_path: Path, **overrides: object) -> MultimodalService:
    settings = SimpleNamespace(
        data_dir=tmp_path,
        video_max_keyframes=24,
        video_min_frame_spacing_seconds=2.0,
        ffmpeg_timeout_seconds=5.0,
        ffmpeg_threads=1,
        **overrides,
    )
    return MultimodalService(settings, StorageService(settings), SimpleNamespace())  # type: ignore[arg-type]


@pytest.mark.parametrize(
    ("text", "coverage", "ink", "expected"),
    [
        ("文" * 100, 0.9, 0.1, VisualClassification.DOCUMENT),
        ("dashboard values", 0.9, 0.1, VisualClassification.MIXED),
        ("", 0.9, 0.1, VisualClassification.VISUAL),
        ("", 0.1, 0.2, VisualClassification.UNKNOWN),
    ],
)
def test_image_routing_uses_measured_signals(
    tmp_path: Path,
    text: str,
    coverage: float,
    ink: float,
    expected: VisualClassification,
) -> None:
    evidence = PageEvidence(
        page_number=1,
        source_kind=PageSourceKind.MIXED,
        image_coverage_ratio=coverage,
        visual_ink_ratio=ink,
    )
    assert service(tmp_path).classify_image(evidence, text) == expected


@pytest.mark.parametrize(
    ("text", "expected"),
    [
        ("文档正文" * 30, VisualClassification.DOCUMENT),
        ("MIXED DASHBOARD measured values", VisualClassification.MIXED),
        ("", VisualClassification.VISUAL),
        ("短标题", VisualClassification.UNKNOWN),
    ],
)
def test_raster_routing_uses_completed_ocr_when_pdf_metrics_do_not_exist(
    tmp_path: Path, text: str, expected: VisualClassification
) -> None:
    assert service(tmp_path).classify_image(None, text) == expected


@pytest.mark.parametrize(
    ("language", "expected"),
    [
        (
            "zh-CN",
            "Write every natural-language output field in Simplified Chinese and set "
            "the language field to zh-CN. Do not answer in English.",
        ),
        (
            "en",
            "Write every natural-language output field in English and set the "
            "language field to en.",
        ),
        ("auto", "Use the main language visible in the image or sampled frames."),
    ],
)
def test_media_language_instruction_is_shared_by_frames_and_video_summary(
    language: str, expected: str
) -> None:
    assert MultimodalService._language_instruction(language) == expected  # type: ignore[arg-type]


def test_keyframe_selection_is_deterministic_bounded_and_covers_endpoints(
    tmp_path: Path,
) -> None:
    multimodal = service(tmp_path)
    candidates = [
        (float(index) * 3.1, (index % 7) / 7) for index in range(1, 300)
    ] + [(62.0, 0.99), (63.0, 0.20)]
    first = multimodal._select_timestamps(1_800.0, candidates)
    second = multimodal._select_timestamps(1_800.0, list(reversed(candidates)))

    assert first == second
    assert first == sorted(first)
    assert first[0] == 0.0
    assert 1_799.0 < first[-1] < 1_800.0
    assert len(first) <= 24
    assert all(0 <= timestamp <= 1_800.0 for timestamp in first)
    assert all(
        right - left >= 2.0
        for left, right in zip(first, first[1:], strict=False)
    )


def test_short_video_keeps_first_and_last_even_inside_dedup_window(tmp_path: Path) -> None:
    timestamps = service(tmp_path)._select_timestamps(1.0, [(0.4, 0.9)])
    assert timestamps == [0.0, 0.99]


def test_tail_sample_stays_before_last_low_frame_rate_timestamp(tmp_path: Path) -> None:
    timestamps = service(tmp_path)._select_timestamps(6.0, [], frame_rate=12.0)
    assert timestamps[-1] == pytest.approx(5.895833, abs=0.000001)


async def test_embedded_video_relationship_is_never_processed_as_video(
    tmp_path: Path,
) -> None:
    multimodal = service(tmp_path)
    path = Path("tests/fixtures/embedded-video.pptx").resolve()
    extracted = multimodal.extract_office_images(path, PPTX_MIME)

    assert extracted
    assert all(item.mime_type.startswith("image/") for item in extracted)
    assert all(not item.filename.endswith(".mp4") for item in extracted)

    describe = AsyncMock(
        return_value=VisualAnalysis(
            classification=VisualClassification.VISUAL,
            summary="演示文稿中的静态海报图像",
            language="zh-CN",
            model="fixture-vlm",
        )
    )
    multimodal._describe = describe  # type: ignore[method-assign]
    multimodal._run_process = AsyncMock(  # type: ignore[method-assign]
        side_effect=AssertionError("embedded videos must never invoke ffmpeg or ffprobe")
    )
    result = await multimodal.parse_office(
        StoredSource(
            path=path,
            filename=path.name,
            mime_type=PPTX_MIME,
            size_bytes=path.stat().st_size,
            page_count=1,
        ),
        ContentParseOptions(),
        content_id="job_embedded_video",
        progress_callback=None,
        cancel_event=None,
    )

    assert result.parse_result is not None
    parse_result = result.parse_result
    assert parse_result.source.kind == SourceKind.PPTX
    assert parse_result.source.slide_count == 1
    assert parse_result.units[0].unit_type == UnitType.SLIDE
    assert parse_result.units[0].blocks
    assert parse_result.assets
    assert all(asset.kind == AssetKind.IMAGE for asset in parse_result.assets)
    assert all("video" not in warning.code for warning in parse_result.warnings)
    assert multimodal._run_process.await_count == 0  # type: ignore[attr-defined]
    assert describe.await_count == 1


async def test_office_duplicate_images_are_analyzed_once(tmp_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches

    # Repeating the same source image creates two placements but one SHA asset.
    multimodal = service(tmp_path)
    image = Path("tests/fixtures/natural-scene.png").resolve()
    path = tmp_path / "duplicate-images.docx"
    document = Document()
    document.add_heading("Duplicate image fixture", 0)
    document.add_picture(str(image), width=Inches(2))
    document.add_picture(str(image), width=Inches(2))
    document.save(path)
    describe = AsyncMock(
        return_value=VisualAnalysis(
            classification=VisualClassification.VISUAL,
            summary="静态插图",
            language="zh-CN",
            model="fixture-vlm",
        )
    )
    multimodal._describe = describe  # type: ignore[method-assign]
    source = StoredSource(
        path=path,
        filename=path.name,
        mime_type=DOCX_MIME,
        size_bytes=path.stat().st_size,
        page_count=1,
    )
    result = await multimodal.parse_office(
        source,
        ContentParseOptions(),
        content_id="job_docx_image",
        progress_callback=None,
        cancel_event=None,
    )
    assert result.parse_result is not None
    assert len(result.parse_result.assets) == 1
    assert len(result.parse_result.assets[0].locations) == 2
    assert describe.await_count == 1


async def test_office_document_image_uses_child_document_path_without_vlm(
    tmp_path: Path,
) -> None:
    multimodal = service(tmp_path)
    visible_text = "文档图片中的可验证正文" * 12
    child_parser = SimpleNamespace(
        parse=AsyncMock(return_value=SimpleNamespace(plain_text=visible_text))
    )
    describe = AsyncMock(side_effect=AssertionError("document image must not invoke VLM"))
    multimodal._describe = describe  # type: ignore[method-assign]
    path = Path("tests/fixtures/embedded-image.docx").resolve()

    result = await multimodal.parse_office(
        StoredSource(
            path=path,
            filename=path.name,
            mime_type=DOCX_MIME,
            size_bytes=path.stat().st_size,
            page_count=1,
        ),
        ContentParseOptions(),
        content_id="job_document_image",
        progress_callback=None,
        cancel_event=None,
        image_parser=child_parser,  # type: ignore[arg-type]
    )

    assert result.parse_result is not None
    analysis = result.parse_result.assets[0].visual_analysis
    assert analysis is not None
    assert analysis.classification == VisualClassification.DOCUMENT
    assert analysis.visible_text == visible_text
    assert visible_text in result.parse_result.renderings.plain_text
    assert child_parser.parse.await_count == 1
    assert describe.await_count == 0


async def test_embedded_image_vlm_failure_returns_partial_with_original_asset(
    tmp_path: Path,
) -> None:
    multimodal = service(tmp_path)
    multimodal._describe = AsyncMock(  # type: ignore[method-assign]
        side_effect=ParserUnavailableError("fixture VLM unavailable")
    )
    path = Path("tests/fixtures/embedded-image.docx").resolve()
    result = await multimodal.parse_office(
        StoredSource(
            path=path,
            filename=path.name,
            mime_type=DOCX_MIME,
            size_bytes=path.stat().st_size,
            page_count=1,
        ),
        ContentParseOptions(),
        content_id="job_partial_docx",
        progress_callback=None,
        cancel_event=None,
    )

    assert result.parse_result is not None
    parse_result = result.parse_result
    assert parse_result.status == "partial"
    assert len(parse_result.assets) == 1
    assert parse_result.assets[0].status.value == "failed"
    assert multimodal.storage.asset_path(
        "job_partial_docx", parse_result.assets[0].asset_id
    ) is not None
    assert [warning.code for warning in parse_result.warnings] == [
        "embedded_image_analysis_failed"
    ]


async def test_standalone_video_requires_vlm_before_ffmpeg(tmp_path: Path) -> None:
    multimodal = service(tmp_path)
    multimodal.vlm = SimpleNamespace(enabled=False)  # type: ignore[assignment]
    path = Path("tests/fixtures/scene-switch.mp4").resolve()
    with pytest.raises(ParserUnavailableError, match="standalone video"):
        await multimodal.parse_video(
            StoredSource(
                path=path,
                filename=path.name,
                mime_type="video/mp4",
                size_bytes=path.stat().st_size,
                page_count=1,
            ),
            ContentParseOptions(),
            content_id="job_video_no_vlm",
            progress_callback=None,
            cancel_event=None,
        )


async def test_media_process_honors_cancellation(tmp_path: Path) -> None:
    multimodal = service(tmp_path)
    cancelled = asyncio.Event()
    task = asyncio.create_task(
        multimodal._run_process(
            sys.executable,
            "-c",
            "import time; time.sleep(10)",
            cancel_event=cancelled,
        )
    )
    await asyncio.sleep(0.05)
    cancelled.set()
    with pytest.raises(ParserCancelledError):
        await asyncio.wait_for(task, timeout=2)


async def test_missing_media_tool_is_reported_as_backend_unavailable(tmp_path: Path) -> None:
    with pytest.raises(ParserUnavailableError):
        await service(tmp_path)._run_process("definitely-missing-mosaicparse-media-tool")


async def test_tiff_keeps_original_bytes_and_adds_png_preview(tmp_path: Path) -> None:
    multimodal = service(tmp_path)
    content = Path("tests/fixtures/natural-scene.tiff").read_bytes()
    location = AssetLocation(unit_id="unit-1", page_number=1)
    original, _ = await multimodal._persist_image_asset(
        "job_tiff",
        content,
        filename="natural-scene.tiff",
        role=AssetRole.SOURCE,
        location=location,
    )
    preview = await multimodal._preview_asset(
        "job_tiff", original, content, location
    )

    assert original.mime_type == "image/tiff"
    assert multimodal.storage.asset_path("job_tiff", original.asset_id).read_bytes() == content  # type: ignore[union-attr]
    assert preview is not None
    assert preview.role == AssetRole.PREVIEW
    assert preview.mime_type == "image/png"
    assert preview.parent_asset_id == original.asset_id
    preview_path = multimodal.storage.asset_path("job_tiff", preview.asset_id)
    assert preview_path is not None
    assert preview_path.parent.name == "previews"


@pytest.mark.skipif(
    shutil.which("ffmpeg") is None or shutil.which("ffprobe") is None,
    reason="real FFmpeg acceptance runs in the application image",
)
async def test_real_ffmpeg_video_pipeline_produces_bounded_keyframe_evidence(
    tmp_path: Path,
) -> None:
    class FakeVlm:
        enabled = True
        name = "fixture-vlm"
        model = "fixture-vlm"

        async def complete_structured(
            self, images, prompt, response_model, **kwargs
        ):  # type: ignore[no-untyped-def]
            if "classification" in response_model.model_fields:
                value = response_model(
                    classification="visual",
                    summary=f"sampled frame with {len(images)} image",
                    detailed_description="solid-color sampled test frame",
                    language="en",
                )
            else:
                value = response_model(
                    summary="Summary based only on the sampled test frames."
                )
            return SimpleNamespace(value=value, duration_ms=1)

    settings = SimpleNamespace(
        data_dir=tmp_path,
        video_max_keyframes=6,
        video_min_frame_spacing_seconds=2.0,
        ffmpeg_timeout_seconds=30.0,
        ffmpeg_threads=1,
        ffmpeg_max_alloc_bytes=128 * 1024 * 1024,
    )
    multimodal = MultimodalService(
        settings, StorageService(settings), FakeVlm()  # type: ignore[arg-type]
    )
    path = Path("tests/fixtures/scene-switch.mp4").resolve()
    progress: list[tuple[int, int, str]] = []
    result = await multimodal.parse_video(
        StoredSource(
            path=path,
            filename=path.name,
            mime_type="video/mp4",
            size_bytes=path.stat().st_size,
            page_count=1,
        ),
        ContentParseOptions(description_language="en"),
        content_id="job_real_video",
        progress_callback=lambda current, total, state: progress.append(
            (current, total, state)
        ),
        cancel_event=None,
    )

    assert result.parse_result is not None
    parse_result = result.parse_result
    assert parse_result.video_analysis is not None
    video = parse_result.video_analysis
    timestamps = [frame.timestamp_ms for frame in video.keyframes]
    assert 1 <= len(timestamps) <= 6
    assert timestamps == sorted(timestamps)
    assert timestamps[0] == 0
    assert timestamps[-1] <= video.duration_ms
    assert video.summary == "Summary based only on the sampled test frames."
    assert all(state == "frame.processed" for _, _, state in progress)
    assert progress[-1][:2] == (len(timestamps), len(timestamps))
    for asset in parse_result.assets:
        stored = multimodal.storage.asset_path("job_real_video", asset.asset_id)
        assert stored is not None
        assert hashlib.sha256(stored.read_bytes()).hexdigest() == asset.sha256
